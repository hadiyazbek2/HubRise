"""
Generates HubRise Senior Project Report as a .docx file.
Font: Arial, Body: 12pt, Titles: 16pt, Page numbers: bottom-right.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p

def subheading(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F2F2F2")
    p._p.pPr.append(shading) if p._p.pPr is not None else None
    return p

def page_break():
    doc.add_page_break()

def add_page_numbers():
    """Add page numbers to bottom-right of all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.font.name = "Arial"
        run.font.size = Pt(10)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

# ── 1. COVER PAGE ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Cm(4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("HubRise")
run.font.name = "Arial"
run.font.size = Pt(28)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A Goal-Oriented Social Network for Android")
run.font.name = "Arial"
run.font.size = Pt(18)
run.font.bold = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Senior Project Report")
run.font.name = "Arial"
run.font.size = Pt(16)

doc.add_paragraph()

for label, value in [
    ("Student Name:", "Hadi Yazbek"),
    ("Supervisor:", "Dr. Mohamad Rida Mortada"),
    ("Academic Year:", "2025 – 2026"),
    ("Department:", "Computer Science, Engineering"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label}  ")
    r1.font.name = "Arial"
    r1.font.size = Pt(13)
    r1.font.bold = True
    r2 = p.add_run(value)
    r2.font.name = "Arial"
    r2.font.size = Pt(13)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GitHub Repository: https://github.com/hadiyazbek2/HubRise")
run.font.name = "Arial"
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x00, 0x56, 0xD2)

page_break()

# ── 2. TABLE OF CONTENTS (manual) ─────────────────────────────────────────────
heading("Table of Contents")
toc_entries = [
    ("1. Abstract", 3),
    ("2. Introduction", 3),
    ("3. Related Work", 4),
    ("4. Proposed Project", 5),
    ("5. Project Description", 6),
    ("   5.1 User Registration and Authentication", 6),
    ("   5.2 Home Feed", 7),
    ("   5.3 Hubs", 7),
    ("   5.4 Challenges", 8),
    ("   5.5 Explore Screen", 9),
    ("   5.6 Profile and Social Features", 9),
    ("   5.7 Notifications", 10),
    ("   5.8 Search", 10),
    ("   5.9 Comments", 11),
    ("6. Code Description", 11),
    ("   6.1 RetrofitClient (Kotlin)", 11),
    ("   6.2 UserPreferences (Kotlin)", 12),
    ("   6.3 Challenge Model (Python/Django)", 13),
    ("   6.4 build_auth_response (Python/Django)", 14),
    ("7. Limitations of the Application", 15),
    ("8. Future Work and Conclusion", 15),
    ("9. References", 16),
    ("10. GitHub Repository", 16),
]
for text, _ in toc_entries:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(3)

page_break()

# ── 3. ABSTRACT ───────────────────────────────────────────────────────────────
heading("1. Abstract")
body(
    "HubRise is a goal-oriented social networking application developed for the Android platform. "
    "Unlike conventional social networks where users share arbitrary content, HubRise is built around "
    "the concept of goal-driven communities called Hubs. Each Hub is centered on a shared objective, "
    "and members collaborate, motivate each other, and track real progress together. The application "
    "combines elements of social networking, goal tracking, and community accountability to create a "
    "unique and motivating digital environment."
)
body(
    "The system is built using a modern client-server architecture. The backend is powered by Django "
    "and Django REST Framework with JWT-based authentication, while the Android client is developed in "
    "Kotlin following the MVVM architectural pattern. Data is exchanged over a REST API, and user "
    "sessions are managed securely using JSON Web Tokens stored in Android DataStore, never in "
    "SharedPreferences, to minimize security risks."
)
body(
    "HubRise supports three distinct challenge models: Stage-Based challenges where members progress "
    "through ordered milestones; Count-Based challenges where members log numeric entries toward a "
    "target; and Streak-Based challenges where members check in daily or weekly over a defined period. "
    "Each challenge submission is subject to a peer micro-validation layer, where other Hub members "
    "can mark progress as trusted, adding social accountability to personal goal achievement."
)
body(
    "Additional features include a real-time notification system for follows, likes, comments, and "
    "challenge completion events; a global and personalized post feed; a search feature covering "
    "users, hubs, posts, and challenges; an Explore screen for video and media content; and a "
    "full profile system with follow/following relationships. The project demonstrates the practical "
    "integration of a production-grade REST API with a fully functional Android mobile application."
)

page_break()

# ── 4. INTRODUCTION ───────────────────────────────────────────────────────────
heading("2. Introduction")
body(
    "Social media applications have become an integral part of daily life, yet most existing platforms "
    "focus on passive content consumption rather than active personal growth. Users scroll through "
    "feeds of entertainment without a clear structure to help them achieve meaningful goals. HubRise "
    "addresses this gap by designing a social network where every interaction is tied to a goal."
)
body(
    "The core unit of HubRise is the Hub — a community that exists for a single shared purpose. "
    "Whether the goal is running 100 km, reading 12 books in a year, or building a daily meditation "
    "habit, a Hub provides the structure, the audience, and the accountability needed to succeed. "
    "Members post progress updates, complete challenges, and receive validation from their peers."
)
body(
    "From a technical perspective, HubRise is a full-stack mobile application. The backend exposes a "
    "comprehensive RESTful API with over 25 endpoints covering authentication, community management, "
    "challenge tracking, social interactions, and notifications. The Android application consumes "
    "this API and presents a polished, responsive user interface designed for smooth everyday use."
)
body(
    "This report documents the design decisions, architecture, features, code implementation, and "
    "future direction of the HubRise project. It also discusses the limitations encountered during "
    "development and the improvements planned for subsequent versions."
)

page_break()

# ── 5. RELATED WORK ───────────────────────────────────────────────────────────
heading("3. Related Work")
body(
    "Several existing applications share aspects of HubRise's vision, but none combine all its "
    "elements into a single cohesive product. This section reviews the most closely related systems."
)

subheading("3.1 Strava")
body(
    "Strava is a fitness tracking social network primarily focused on running and cycling. It allows "
    "users to log workouts, follow friends, and join challenges. However, Strava is limited to "
    "physical fitness and does not support custom goal types, community-created challenges, or "
    "non-athletic pursuits such as reading or learning."
)

subheading("3.2 Habitica")
body(
    "Habitica gamifies habit tracking by turning daily tasks into a role-playing game. While it "
    "successfully motivates users through game mechanics, its social layer is shallow — users "
    "interact within parties but there is no rich community feed, no media sharing, and no "
    "peer validation of progress."
)

subheading("3.3 Bereal / Instagram")
body(
    "General social networks like Instagram and BeReal capture spontaneous moments and drive "
    "engagement through likes and comments, but they lack any goal-tracking infrastructure. "
    "Users have no accountability mechanism, no challenge system, and no concept of shared "
    "objective-driven communities."
)

subheading("3.4 HubRise Contributions")
body(
    "HubRise improves upon these applications in several key ways: it supports any goal type "
    "through its three flexible challenge models (stage, count, streak); it combines social "
    "networking features (feed, likes, comments, follow/following) with structured accountability; "
    "it introduces peer micro-validation so community members can vouch for each other's progress; "
    "and it provides Hub-level roles (admin, moderator, member) for community governance. No "
    "existing application delivers all of these features in a single mobile experience."
)

page_break()

# ── 6. PROPOSED PROJECT ───────────────────────────────────────────────────────
heading("4. Proposed Project")
body(
    "HubRise proposes a goal-oriented social network that fills the gaps identified in existing "
    "applications. The core proposition is: give users a meaningful reason to open a social app "
    "every day — not to consume entertainment, but to make progress on something that matters to them."
)
body(
    "The proposed solution is built on three pillars:"
)
bullet(
    "Structured Communities (Hubs): Instead of following individuals, users join goal-specific "
    "communities. Every Hub has a defined purpose, and all content posted within it is related "
    "to that purpose."
)
bullet(
    "Flexible Challenge System: Hubs can run challenges in three models — Stage-Based (sequential "
    "milestones), Count-Based (numeric targets), and Streak-Based (daily/weekly consistency). "
    "This covers virtually any goal type imaginable."
)
bullet(
    "Social Accountability Layer: Progress posts can be validated by community members ("
    "peer micro-validation). Challenge completions require admin approval, ensuring that "
    "achievements are genuine. Notifications keep the community engaged and supportive."
)
body(
    "By solving the limitation of existing apps — either too narrow in goal type (Strava), "
    "too gamified without community (Habitica), or lacking goal structure entirely (Instagram) — "
    "HubRise provides a complete environment where social motivation and personal achievement "
    "reinforce each other."
)

page_break()

# ── 7. PROJECT DESCRIPTION ────────────────────────────────────────────────────
heading("5. Project Description")
body(
    "This section describes each feature of HubRise in detail. The application is organized into "
    "a Django REST backend and an Android Kotlin frontend. The following subsections cover each "
    "major functional area."
)

subheading("5.1 User Registration and Authentication")
body(
    "New users register through a three-step onboarding flow. In Step 1, they provide an email "
    "address and password. The app validates the email format and checks availability in real time "
    "by calling the /accounts/check-email/ endpoint. In Step 2, users choose a unique username "
    "and enter their full name and date of birth. In Step 3, users select interests from a "
    "predefined list and optionally upload a profile picture."
)
body(
    "Authentication uses JWT tokens. On login, the backend returns an access token (30-minute TTL) "
    "and a refresh token (7-day TTL). Both are stored securely in Android DataStore. The Retrofit "
    "HTTP client automatically attaches the access token to every request via an OkHttp Interceptor, "
    "and a separate OkHttp Authenticator silently refreshes the token on 401 errors without "
    "interrupting the user experience."
)
body(
    "On logout, the refresh token is blacklisted server-side via the /accounts/logout/ endpoint, "
    "ensuring the session cannot be reused. Social login (Google, Facebook, Apple) is also "
    "supported at the API level."
)

subheading("5.2 Home Feed")
body(
    "The Home screen displays a personalized feed of posts from Hubs the user has joined. Each "
    "post card shows the author's avatar, username, the Hub it belongs to, the post content, any "
    "attached media, a like button with live count, a comment button, and a timestamp. Users can "
    "like and unlike posts with a single tap; the UI updates immediately via optimistic state "
    "management in the HomeViewModel. A global feed tab shows all public posts from all Hubs."
)

subheading("5.3 Hubs")
body(
    "The Hubs screen is divided into two tabs via a ViewPager2: My Hubs (communities the user "
    "has joined) and Explore Hubs (all public Hubs with a recommended section). Each Hub card "
    "shows the cover image, name, description, member count, and category."
)
body(
    "Tapping a Hub opens the Hub Detail screen, which shows the Hub's posts, its active challenges, "
    "a member list, and — for admins — a settings button to edit Hub details or manage members. "
    "Users can join or leave a Hub with a single button. The creator of a Hub automatically "
    "becomes its admin."
)
body(
    "Creating a Hub requires a name, description, optional cover image, and a category (interest "
    "tag). The new Hub is immediately visible in the Explore Hubs list and on the creator's "
    "My Hubs tab."
)

subheading("5.4 Challenges")
body(
    "Challenges are the core engagement mechanism of HubRise. Each Hub can have multiple challenges. "
    "When creating a challenge, the Hub admin selects one of three progress models:"
)
bullet(
    "Stage-Based: The challenge is divided into ordered stages. Each stage can require a specific "
    "type of proof (photo, video, text, number, or any). Members complete stages sequentially. "
    "A visual progress bar tracks how many stages have been completed."
)
bullet(
    "Count-Based: The challenge has a numeric target (e.g., run 100 km, read 12 books). Members "
    "log entries that increment toward the target. The admin can configure the unit label, "
    "increment size, and whether entries are additive or cumulative."
)
bullet(
    "Streak-Based: Members must check in at a defined frequency (daily or weekly) for a target "
    "number of days. A calendar heatmap visualizes check-in history. Grace days can be configured "
    "to allow occasional misses without breaking the streak."
)
body(
    "Challenge templates are available to speed up creation. Official templates (e.g., '30-Day "
    "Fitness Challenge', 'Read 12 Books This Year') come pre-loaded. When a member believes they "
    "have completed a challenge, they submit a Completion Request with an optional note. The Hub "
    "admin reviews it and approves or rejects it. On approval, an achievement broadcast post "
    "is automatically published to the Hub feed."
)

subheading("5.5 Explore Screen")
body(
    "The Explore screen presents public posts from across all Hubs in a vertically scrollable, "
    "full-screen format optimized for video and image content. Posts are fetched paginated "
    "via the /community/posts/explore/ endpoint and displayed using a RecyclerView with "
    "smooth auto-play for video content."
)

subheading("5.6 Profile and Social Features")
body(
    "Each user has a profile page showing their avatar, full name, bio, post count, Hub count, "
    "follower count, and following count. Users can follow and unfollow each other. Following "
    "a user generates a notification. The profile page also lists all posts by that user "
    "that the viewer has permission to see (posts in public Hubs are always visible; posts "
    "in private Hubs are only visible to members)."
)
body(
    "Profile editing allows users to update their full name, bio, and wishlist URL. The wishlist "
    "URL is surfaced on achievement posts to allow supporters to send gifts to challenge completers."
)

subheading("5.7 Notifications")
body(
    "The Notifications screen displays a reverse-chronological list of events relevant to the "
    "current user: new followers, likes on their posts, comments on their posts, challenge "
    "completion request submissions (for Hub admins), and challenge completion approvals or "
    "rejections (for challenge participants). All unread notifications are automatically marked "
    "as read when the screen is opened. An unread badge count is shown on the notification "
    "tab icon."
)

subheading("5.8 Search")
body(
    "The Search screen allows users to search across four categories simultaneously: Users, "
    "Hubs, Posts, and Challenges. Results are returned from the /community/search/ endpoint "
    "and displayed in tab-separated sections. Each result type has a dedicated adapter with "
    "an appropriate card layout."
)

subheading("5.9 Comments")
body(
    "Users can comment on any post. Tapping the comment button on a post card opens a "
    "BottomSheetFragment that displays existing comments and provides an input field for new "
    "ones. Comments are loaded from the /community/posts/{id}/comments/ endpoint and submitted "
    "via POST. Comment counts are shown on the post card."
)

page_break()

# ── 8. CODE DESCRIPTION ───────────────────────────────────────────────────────
heading("6. Code Description")
body(
    "This section documents the most architecturally significant classes and methods in the "
    "HubRise codebase. These were selected because they implement non-trivial logic that "
    "underpins the security, networking, and data modeling of the entire application."
)

subheading("6.1 RetrofitClient.kt — Token Interceptor and Authenticator")
body(
    "File: app/src/main/java/com/example/hubrise/data/api/RetrofitClient.kt"
)
body(
    "RetrofitClient is a Kotlin singleton object that configures and provides all Retrofit "
    "instances used by the application. It solves two critical problems: attaching the JWT "
    "access token to every outgoing request, and automatically refreshing the token when it "
    "expires (HTTP 401 response)."
)
body("The token interceptor reads the access token from DataStore and injects it as a Bearer header:")
code_block(
    'private val tokenInterceptor = Interceptor { chain ->\n'
    '    val prefs = UserPreferences(appContext)\n'
    '    val token = runBlocking { prefs.accessToken.first() }\n'
    '    val request = if (!token.isNullOrEmpty()) {\n'
    '        chain.request().newBuilder()\n'
    '            .header("Authorization", "Bearer $token")\n'
    '            .build()\n'
    '    } else { chain.request() }\n'
    '    chain.proceed(request)\n'
    '}'
)
body(
    "The tokenAuthenticator implements OkHttp's Authenticator interface and is invoked "
    "automatically whenever a 401 Unauthorized response is received. It performs a synchronous "
    "token refresh call using a separate no-auth Retrofit instance (to avoid triggering the "
    "interceptor again), saves the new access token, and retries the original request. If the "
    "refresh itself fails, it clears all stored credentials and redirects the user to the "
    "LoginActivity. This design means the user is never shown an error for an expired token — "
    "the refresh is completely transparent."
)
body(
    "RetrofitClient also detects whether the app is running on an emulator (via Build fingerprint "
    "checks) and automatically selects the correct base URL (10.0.2.2 for emulator, LAN IP for "
    "physical device)."
)

subheading("6.2 UserPreferences.kt — Secure Token Storage")
body(
    "File: app/src/main/java/com/example/hubrise/data/local/UserPreferences.kt"
)
body(
    "UserPreferences wraps Android's Jetpack DataStore (Preferences) to provide a type-safe, "
    "coroutine-friendly interface for all persistent user session data. Unlike SharedPreferences, "
    "DataStore is backed by Protocol Buffers and provides atomic read-write guarantees, making "
    "it significantly more reliable for concurrent access from coroutines."
)
body("Each piece of data has a typed key and a corresponding Flow for reactive observation:")
code_block(
    'private val ACCESS_TOKEN = stringPreferencesKey("access_token")\n\n'
    'val accessToken: Flow<String?> = context.dataStore.data.map { preferences ->\n'
    '    preferences[ACCESS_TOKEN]\n'
    '}\n\n'
    'suspend fun saveAccessToken(token: String) {\n'
    '    context.dataStore.edit { preferences ->\n'
    '        preferences[ACCESS_TOKEN] = token\n'
    '    }\n'
    '}'
)
body(
    "The class stores access token, refresh token, user ID, email, username, full name, "
    "profile picture URL, login status, and last login timestamp. The clearAll() method "
    "atomically wipes the entire DataStore on logout. The hasToken() suspend function is "
    "used at app startup to decide whether to show the login screen or navigate directly "
    "to the main activity."
)

subheading("6.3 Challenge Model — Django (community/models.py)")
body(
    "File: backend/community/models.py"
)
body(
    "The Challenge model is the centerpiece of the HubRise data layer. It uses a single "
    "polymorphic design: one Challenge row with a progress_model field ('stage', 'count', or "
    "'streak') acts as the header, and separate related models (ChallengeStage, "
    "ChallengeCountConfig, ChallengeStreakConfig) store the model-specific configuration. "
    "This avoids nullable columns on a single fat model and keeps each configuration clean."
)
code_block(
    'class Challenge(models.Model):\n'
    '    MODEL_STAGE = "stage"\n'
    '    MODEL_COUNT = "count"\n'
    '    MODEL_STREAK = "streak"\n\n'
    '    hub = models.ForeignKey(Hub, on_delete=models.CASCADE, related_name="challenges")\n'
    '    title = models.CharField(max_length=100)\n'
    '    progress_model = models.CharField(max_length=10, choices=MODEL_CHOICES)\n'
    '    is_main = models.BooleanField(default=False)\n'
    '    ends_at = models.DateTimeField(null=True, blank=True)'
)
body(
    "User progress is tracked in separate rows per user per challenge: UserStageProgress "
    "(stage completion status and proof post link), UserCountProgress (running total and "
    "completion flag), and UserStreakProgress (current streak, longest streak, total check-ins, "
    "and a JSON calendar array for the heatmap visualization). This design supports thousands "
    "of concurrent participants per challenge without schema changes."
)

subheading("6.4 build_auth_response — Django (accounts/views.py)")
body(
    "File: backend/accounts/views.py"
)
body(
    "The build_auth_response helper function is called by SignupView, LoginView, and "
    "SocialLoginView to generate a consistent authentication response. It uses "
    "djangorestframework-simplejwt to generate a JWT refresh token (from which the access token "
    "is derived), then serializes the user object with the UserSerializer."
)
code_block(
    'def build_auth_response(user, request):\n'
    '    refresh = RefreshToken.for_user(user)\n'
    '    return {\n'
    '        "access_token": str(refresh.access_token),\n'
    '        "refresh_token": str(refresh),\n'
    '        "user": UserSerializer(user, context={"request": request}).data,\n'
    '    }'
)
body(
    "By centralizing the token generation logic in a single function, all three authentication "
    "entry points return exactly the same response structure. The Android client can therefore "
    "use a single AuthResponse data class to parse all three flows. The request context is "
    "passed to the serializer so that profile_picture_url is returned as an absolute URL "
    "(using request.build_absolute_uri) rather than a relative path."
)

page_break()

# ── 9. LIMITATIONS ────────────────────────────────────────────────────────────
heading("7. Limitations of the Application")
body(
    "Despite its comprehensive feature set, HubRise has several known limitations in its "
    "current version:"
)
bullet(
    "Database: The application currently uses SQLite for development. SQLite is not suitable "
    "for production workloads with concurrent writes. Migration to PostgreSQL is planned but "
    "not yet completed."
)
bullet(
    "Social Login: Google, Facebook, and Apple social login endpoints exist in the API but "
    "do not perform cryptographic signature verification of the identity tokens. This means "
    "social login is a prototype and should not be used in a production environment without "
    "proper token verification."
)
bullet(
    "Push Notifications: The current notification system is pull-based — the client fetches "
    "notifications on demand. There are no push notifications (FCM/APNs), so users are not "
    "alerted in real time when they are offline."
)
bullet(
    "Media Storage: Uploaded images and videos are stored on the local filesystem of the "
    "Django server. In a production deployment, these should be stored in a cloud object "
    "store (e.g., AWS S3 or Google Cloud Storage) with a CDN for delivery."
)
bullet(
    "Scalability: The backend runs as a single Django process with no caching layer (Redis), "
    "no task queue (Celery), and no horizontal scaling. High traffic would require these "
    "additions."
)
bullet(
    "No End-to-End Tests: The project has unit test stubs but no comprehensive test suite. "
    "Full end-to-end and integration tests are needed before production deployment."
)

page_break()

# ── 10. FUTURE WORK AND CONCLUSION ────────────────────────────────────────────
heading("8. Future Work and Conclusion")
subheading("Future Work")
body(
    "The following enhancements are planned for future iterations of HubRise:"
)
bullet(
    "Push Notifications via Firebase Cloud Messaging (FCM) to deliver real-time alerts to "
    "users even when the app is in the background."
)
bullet(
    "PostgreSQL migration for the backend database to support concurrent production workloads."
)
bullet(
    "AWS S3 or Google Cloud Storage integration for scalable media file storage and "
    "CDN-backed delivery."
)
bullet(
    "Achievement Badges: A visual badge system where completing challenges earns permanent "
    "badges displayed on the user's profile."
)
bullet(
    "Leaderboards: Hub-level leaderboards ranking members by challenge progress, "
    "encouraging friendly competition."
)
bullet(
    "Direct Messaging: Private messaging between Hub members to enable off-feed coordination."
)
bullet(
    "AI-Powered Challenge Suggestions: Using user interests and Hub data to suggest "
    "relevant challenges to join."
)
bullet(
    "iOS Version: A Swift/SwiftUI client consuming the same REST API to reach Apple "
    "device users."
)

subheading("Conclusion")
body(
    "HubRise successfully demonstrates that a social network can be built around goals rather "
    "than entertainment. The project delivers a fully functional Android application backed by "
    "a production-quality REST API, with features spanning authentication, community management, "
    "three distinct challenge tracking models, peer validation, social interactions, notifications, "
    "and search."
)
body(
    "The architecture — Django REST Framework on the backend, Kotlin MVVM with Retrofit and "
    "DataStore on Android — follows industry best practices and provides a solid foundation "
    "for future development. The peer micro-validation layer and the three-tier challenge "
    "completion review process (self-submission → peer validation → admin approval) are "
    "original contributions that distinguish HubRise from all existing applications surveyed "
    "in the Related Work section."
)
body(
    "This project represents a complete senior-year capstone that integrates backend engineering, "
    "mobile development, database design, RESTful API design, and UI/UX considerations into "
    "a single cohesive product. It is ready for demonstration as a working prototype and "
    "serves as the foundation for a production-ready application."
)

page_break()

# ── 11. REFERENCES ────────────────────────────────────────────────────────────
heading("9. References")
refs = [
    "Django Software Foundation. (2024). Django Documentation. https://docs.djangoproject.com/",
    "Django REST Framework. (2024). DRF Documentation. https://www.django-rest-framework.org/",
    "SimpleJWT. (2024). Django REST Framework SimpleJWT. https://django-rest-framework-simplejwt.readthedocs.io/",
    "Google. (2024). Android Kotlin Documentation. https://developer.android.com/kotlin",
    "Square. (2024). Retrofit HTTP Client. https://square.github.io/retrofit/",
    "Google. (2024). Android Jetpack DataStore. https://developer.android.com/topic/libraries/architecture/datastore",
    "Google. (2024). Android Navigation Component. https://developer.android.com/guide/navigation",
    "Strava. (2024). Strava Platform Documentation. https://developers.strava.com/",
    "Habitica. (2024). Habitica — Gamify Your Life. https://habitica.com/",
    "OkHttp. (2024). OkHttp Documentation. https://square.github.io/okhttp/",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"[{i}]  {ref}")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.5)

page_break()

# ── 12. GITHUB REPOSITORY ─────────────────────────────────────────────────────
heading("10. GitHub Repository")
body(
    "The complete source code for HubRise (both the Django backend and the Android Kotlin "
    "frontend) is publicly available at the following GitHub repository:"
)
p = doc.add_paragraph()
run = p.add_run("https://github.com/hadiyazbek2/HubRise")
run.font.name = "Arial"
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x56, 0xD2)
p.paragraph_format.space_before = Pt(6)

body(
    "The repository contains the full commit history, all source files for backend and Android, "
    "the database migration history, and this project report. The backend can be run locally "
    "by creating a Python virtual environment, installing requirements, and running "
    "'python manage.py runserver'. The Android project can be opened directly in Android Studio."
)

# ── Page numbers ──────────────────────────────────────────────────────────────
add_page_numbers()

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "/home/hadi/AndroidStudioProjects/HubRise/HubRise_Senior_Project_Report.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
